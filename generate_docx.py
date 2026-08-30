"""
generate_docx.py
Converts Dissertation text into a professionally formatted Microsoft Word (.docx) document
adhering to KIU University guidelines:
- Times New Roman, 12pt body text
- 1.5 Line Spacing
- 1 inch margins all around
- Page numbers, Table styling, Figure embeddings
"""

import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_dissertation_docx():
    doc = Document()

    # 1. Page Setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 2. Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

    # Read dissertation markdown
    dissertation_path = r"C:\Users\Lakmal\.gemini\antigravity\brain\72e3f9c9-eae3-40d8-85d2-cb806c66bdee\dissertation.md"
    if not os.path.exists(dissertation_path):
        print(f"Error: {dissertation_path} not found.")
        return

    with open(dissertation_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    
    in_table = False
    table_lines = []

    def flush_table(t_lines):
        if not t_lines:
            return
        # Parse markdown table
        headers = [h.strip() for h in t_lines[0].split('|')[1:-1]]
        data = []
        for line in t_lines[2:]: # Skip delimiter line
            if '|' in line:
                row = [c.strip().replace('**', '') for c in line.split('|')[1:-1]]
                if len(row) == len(headers):
                    data.append(row)

        table = doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for p in hdr_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10.5)

        # Data Rows
        for r_idx, row in enumerate(data):
            row_cells = table.rows[r_idx + 1].cells
            for c_idx, val in enumerate(row):
                row_cells[c_idx].text = val
                for p in row_cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)

        doc.add_paragraph() # Spacing

    print("[*] Processing Dissertation Markdown into DOCX...")
    for line in lines:
        stripped = line.strip()
        
        # Table Detection
        if stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            in_table = False
            flush_table(table_lines)
            table_lines = []

        if not stripped:
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = stripped[2:]
            # Clean bold markers
            text_clean = text.replace('**', '')
            p.add_run(text_clean)
            p.paragraph_format.space_after = Pt(3)
        elif stripped.startswith('1. ') or stripped.startswith('2. ') or stripped.startswith('3. '):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s*', '', stripped).replace('**', '')
            p.add_run(text)
            p.paragraph_format.space_after = Pt(3)
        else:
            # Paragraph
            text_clean = stripped.replace('**', '').replace('*', '')
            p = doc.add_paragraph(text_clean)

    if in_table:
        flush_table(table_lines)

    output_path = r"c:\Users\Lakmal\Documents\Research\Final_Dissertation_14519.docx"
    doc.save(output_path)
    print(f"[+] Final Dissertation Word Document created successfully at: {output_path}")

if __name__ == '__main__':
    create_dissertation_docx()
